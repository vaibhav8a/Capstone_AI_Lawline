"""
export.py — API router for exporting results as PDF or DOCX (Word)
"""

import io
from typing import List, Optional
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from backend.routers.auth import require_permission

router = APIRouter()


class ExportRequest(BaseModel):
    format: str           # "pdf" or "docx"
    title: str = "Legal Research Export"
    query: str = ""
    answer: str = ""
    chunks: Optional[List[dict]] = []


def _build_docx(req: ExportRequest) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(req.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Query
    doc.add_heading("Research Query", level=1)
    doc.add_paragraph(req.query)

    # Answer
    doc.add_heading("Generated Answer", level=1)
    doc.add_paragraph(req.answer)

    # Sources
    if req.chunks:
        doc.add_heading("Cited Sources", level=1)
        for i, chunk in enumerate(req.chunks, 1):
            doc.add_heading(f"{i}. {chunk.get('case_title', 'Unknown')}", level=2)
            p = doc.add_paragraph()
            p.add_run("Court: ").bold = True
            p.add_run(chunk.get("court", "N/A"))
            p2 = doc.add_paragraph()
            p2.add_run("Date: ").bold = True
            p2.add_run(chunk.get("date", "N/A"))
            p3 = doc.add_paragraph()
            p3.add_run("Section Type: ").bold = True
            p3.add_run(chunk.get("section_type", "N/A"))
            doc.add_paragraph(chunk.get("text", "")[:500] + "...")
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(req: ExportRequest) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle("Title", parent=styles["Title"],
                                 fontSize=18, textColor=colors.HexColor("#3b82f6"),
                                 spaceAfter=12)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"],
                         fontSize=13, textColor=colors.HexColor("#1e293b"), spaceAfter=6)
    body = ParagraphStyle("Body", parent=styles["Normal"],
                          fontSize=10, leading=14, spaceAfter=8)
    label = ParagraphStyle("Label", parent=styles["Normal"],
                           fontSize=9, textColor=colors.HexColor("#64748b"))

    story = [
        Paragraph(req.title, title_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0")),
        Spacer(1, 0.4*cm),
        Paragraph("Research Query", h1),
        Paragraph(req.query or "(no query)", body),
        Spacer(1, 0.3*cm),
        Paragraph("Generated Answer", h1),
        Paragraph((req.answer or "(no answer)").replace("\n", "<br/>"), body),
        Spacer(1, 0.4*cm),
    ]

    if req.chunks:
        story.append(Paragraph("Cited Sources", h1))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e2e8f0")))
        story.append(Spacer(1, 0.2*cm))
        for i, chunk in enumerate(req.chunks, 1):
            story.append(Paragraph(f"{i}. {chunk.get('case_title', 'Unknown')}", h1))
            story.append(Paragraph(
                f"<b>Court:</b> {chunk.get('court', 'N/A')} &nbsp;|&nbsp; "
                f"<b>Date:</b> {chunk.get('date', 'N/A')} &nbsp;|&nbsp; "
                f"<b>Type:</b> {chunk.get('section_type', 'N/A')}", label))
            story.append(Paragraph(chunk.get("text", "")[:500] + "...", body))
            story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    return buf.getvalue()


@router.post("/download")
async def download_export(req: ExportRequest):
    """Exports research results as PDF or DOCX."""
    fmt = req.format.lower()
    if fmt not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="Format must be 'pdf' or 'docx'")

    try:
        if fmt == "pdf":
            data = _build_pdf(req)
            media_type = "application/pdf"
            filename = "legal_research.pdf"
        else:
            data = _build_docx(req)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = "legal_research.docx"

        return StreamingResponse(
            io.BytesIO(data),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except ImportError as e:
        raise HTTPException(status_code=500,
                            detail=f"Missing export library: {e}. Run: pip install python-docx reportlab")
    except Exception as e:
        raise HTTPException(status_code=500, detail="Request failed.")
